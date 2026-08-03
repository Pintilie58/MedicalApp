from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from .helpers import add_heading, add_para, add_bullets, add_table, page_break


def build(doc):
    # -------- COPERTĂ --------
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PLAN DE AFACERI")
    r.font.size = Pt(30)
    r.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x68)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PARTEA A II-A — DESCRIEREA PROIECTULUI")
    r.font.size = Pt(18)
    r.bold = True

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("„MyMedicalApp — Platformă SaaS de interpretare inteligentă "
                  "a analizelor medicale de laborator, bazată pe Inteligență "
                  "Artificială și codificare LOINC\u201d")
    r.font.size = Pt(14)
    r.italic = True

    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Solicitant: S.C. FIXMEDICAL S.R.L.")
    r.font.size = Pt(14)
    r.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Program Operațional Creștere Inteligentă, Digitalizare și "
                  "Instrumente Financiare (POCIDIF) 2021–2027")
    r.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Document elaborat conform structurii din ANEXA 4 — "
                  "Model Plan de Afaceri")
    r.font.size = Pt(11)
    r.italic = True

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Anul redactării: 2026 · Perioada de implementare: 18 luni · "
                  "Buget total al proiectului: 800.000 EUR "
                  "(finanțare nerambursabilă + contribuție proprie)")
    r.font.size = Pt(11)
    page_break(doc)

    # -------- PREAMBUL PARTEA A II-A --------
    add_heading(doc, "PARTEA A II-A — DESCRIEREA PROIECTULUI", 1)
    add_para(doc,
             "Prezenta secțiune descrie în detaliu proiectul de investiții "
             "propus de S.C. FIXMEDICAL S.R.L.: proiectarea, dezvoltarea de la "
             "zero și lansarea comercială a platformei software MyMedicalApp — "
             "o soluție SaaS (Software-as-a-Service) inovatoare care transformă "
             "buletinele de analize medicale de laborator (documente PDF "
             "nestructurate, emise în formate eterogene de sute de laboratoare "
             "diferite) în date medicale structurate, codificate internațional "
             "(LOINC) și interpretate în limbaj accesibil pacientului, în limba "
             "sa maternă, prin tehnologii de Inteligență Artificială generativă.")
    add_para(doc,
             "Proiectul este conceput integral ca o inițiativă nouă "
             "(greenfield): toate componentele software, infrastructura, "
             "procesele operaționale și canalele comerciale vor fi proiectate, "
             "construite, testate și lansate în cadrul perioadei de "
             "implementare de 18 luni, pornind de la stadiul de concept "
             "tehnologic validat teoretic (TRL 2–3) și ajungând la produs "
             "comercial operațional pe piață (TRL 8–9).")
    add_para(doc, "Structura prezentei părți urmează întocmai cerințele "
             "Anexei 4:", bold=True)
    add_bullets(doc, [
        "C.1. Obiectivele proiectului — necesitatea și oportunitatea "
        "investiției, obiectivul general, obiectivele specifice (SMART), "
        "rezultatele așteptate și alinierea strategică;",
        "C.2. Caracterul inovativ al aplicației — stadiul actual al "
        "tehnologiei, elementele de noutate, componenta de Inteligență "
        "Artificială și poziționarea față de soluțiile existente;",
        "C.3. Product-Market Fit — piața țintă, problema abordată, "
        "validarea cererii și propunerea de valoare;",
        "C.4. Descrierea tehnică a proiectului — arhitectura sistemului și "
        "enumerarea completă a etapelor și fazelor de dezvoltare, de la "
        "landing page, înregistrare și autentificare până la motorul AI, "
        "microserviciul de codificare LOINC, modulul B2B și lansarea "
        "comercială;",
        "C.5. Personal și instruire — echipa de proiect, roluri, plan de "
        "formare profesională;",
        "C.6. Graficul estimat al proiectului — calendar, diagramă Gantt, "
        "milestone-uri și dependințe critice;",
        "C.7. Strategia de marketing și de comercializare;",
        "C.8. Impactul proiectului — rezultate cantitative, calitative, "
        "sustenabilitate și contribuție strategică;",
        "C.9. Evaluarea și atenuarea riscurilor."
    ])

    add_para(doc, "Date de identificare sintetice ale proiectului:", bold=True)
    add_table(doc,
              headers=["Element", "Valoare"],
              rows=[
                  ["Denumirea proiectului",
                   "MyMedicalApp — Platformă SaaS de interpretare inteligentă "
                   "a analizelor medicale"],
                  ["Solicitant", "S.C. FIXMEDICAL S.R.L."],
                  ["Program de finanțare",
                   "POCIDIF 2021–2027 — digitalizare și creștere inteligentă"],
                  ["Buget total estimat", "800.000 EUR"],
                  ["Perioada de implementare", "18 luni"],
                  ["Domeniu", "e-Health / HealthTech — software medical (SaaS)"],
                  ["Nivel de maturitate tehnologică",
                   "TRL 2–3 la debut → TRL 8–9 la finalizare"],
                  ["Modele de venit", "B2C freemium/premium · B2B SaaS pe "
                   "abonament · B2B enterprise on-premise"],
                  ["Piețe țintă", "România (lansare), Uniunea Europeană "
                   "(extindere), acoperire lingvistică 7 limbi la lansare, "
                   "30 de limbi la finalul proiectului"],
              ],
              col_widths_cm=[5.5, 10.7])
    page_break(doc)
