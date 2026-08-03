from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _set_cell_bg(cell, color_hex: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_heading(doc, text: str, level: int = 1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x68)
    return h


def add_para(doc, text: str, *, bold=False, italic=False, size=11, justify=True, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    return p


def add_bullets(doc, items, *, size=11):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = p.add_run(it)
        r.font.size = Pt(size)


def add_numbered(doc, items, *, size=11):
    for it in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(2)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = p.add_run(it)
        r.font.size = Pt(size)


def add_table(doc, headers, rows, *, col_widths_cm=None, font_size=10):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Light Grid Accent 1"
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        r.font.bold = True
        r.font.size = Pt(font_size)
        _set_cell_bg(hdr[i], "1F3A68")
        for run in p.runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            if ci >= len(tbl.rows[ri].cells):
                break
            tbl.rows[ri].cells[ci].text = ""
            p = tbl.rows[ri].cells[ci].paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(font_size)
    if col_widths_cm:
        for row in tbl.rows:
            for i, w in enumerate(col_widths_cm):
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)
    return tbl


def page_break(doc):
    doc.add_page_break()


def make_doc():
    doc = Document()
    sec = doc.sections[0]
    sec.left_margin = Cm(2.2)
    sec.right_margin = Cm(2.2)
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    return doc
